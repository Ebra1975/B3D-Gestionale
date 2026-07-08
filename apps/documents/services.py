import re
import shutil
import subprocess
import tempfile
import zipfile
from decimal import Decimal
from pathlib import Path

from django.conf import settings
from django.db import IntegrityError, transaction
from django.db.models import Max
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import DocumentProfile, DocumentTemplate, GeneratedDocument


SUPPORTED_TEMPLATE_VARIABLES = {
    DocumentTemplate.DocumentType.CONSULTING_ESTIMATE: {
        "cliente",
        "preventivo",
        "configurazione",
        "proposta",
        "b3d",
    },
    DocumentTemplate.DocumentType.INTERNAL_ESTIMATE: {
        "cliente",
        "preventivo",
        "configurazione",
        "proposta",
        "b3d",
        "interno",
        "voci_costo",
    },
    DocumentTemplate.DocumentType.SUPPLY_ESTIMATE: {
        "cliente",
        "preventivo",
        "configurazione",
        "proposta",
        "b3d",
        "fornitura",
        "voci_fornitura",
    },
}


def money(value):
    amount = value or Decimal("0.00")
    return f"{amount:.2f} EUR"


def safe_filename(value):
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")
    return cleaned or "documento"


def render_docx_template(template_path, output_path, context):
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = DocxTemplate(template_path)
    document.render(context)
    document.save(output)
    return output


def validate_docx_template_file(uploaded_file, document_type):
    errors = []
    suffix = Path(uploaded_file.name).suffix or ".docx"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temporary_file:
        temporary_path = Path(temporary_file.name)
        for chunk in uploaded_file.chunks():
            temporary_file.write(chunk)

    try:
        if not zipfile.is_zipfile(temporary_path):
            return ["Il file non e un DOCX valido: aprirlo in Word/LibreOffice e salvarlo di nuovo come .docx."]

        try:
            document = DocxTemplate(temporary_path)
        except Exception:
            return ["Il template DOCX non puo essere letto dal gestionale."]

        if hasattr(document, "get_undeclared_template_variables"):
            try:
                variables = set(document.get_undeclared_template_variables())
            except Exception as exc:
                return [f"I segnaposto del template non sono leggibili: {exc}"]

            allowed_variables = SUPPORTED_TEMPLATE_VARIABLES.get(document_type)
            if allowed_variables:
                unknown_variables = sorted(variables - allowed_variables)
                if unknown_variables:
                    errors.append(
                        "Il template contiene segnaposto non riconosciuti: "
                        f"{', '.join(unknown_variables)}."
                    )
    finally:
        temporary_path.unlink(missing_ok=True)
        uploaded_file.seek(0)

    return errors


def find_libreoffice_command():
    configured_path = getattr(settings, "LIBREOFFICE_PATH", "")
    if configured_path and Path(configured_path).exists():
        return configured_path

    windows_path = Path("C:/Program Files/LibreOffice/program/soffice.com")
    if windows_path.exists():
        return str(windows_path)

    return shutil.which("soffice") or shutil.which("libreoffice")


def convert_docx_to_pdf(docx_path, output_dir):
    command = find_libreoffice_command()
    if not command:
        return None

    docx_path = Path(docx_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="b3dlab_lo_profile_") as profile_dir:
        user_installation = Path(profile_dir).resolve().as_uri()
        try:
            result = subprocess.run(
                [
                    command,
                    f"-env:UserInstallation={user_installation}",
                    "--headless",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path),
                ],
                check=False,
                capture_output=True,
                text=True,
                timeout=getattr(settings, "LIBREOFFICE_TIMEOUT_SECONDS", 30),
            )
        except subprocess.TimeoutExpired:
            return None

    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf_path.exists():
        return None
    return pdf_path


def selected_configuration(estimate):
    selected = [configuration for configuration in estimate.configurations.all() if configuration.is_selected]
    if selected:
        return selected[0]
    configurations = list(estimate.configurations.all())
    return configurations[0] if configurations else None


def build_document_export_checks(estimate):
    configuration = selected_configuration(estimate)
    profile = DocumentProfile.get_active()
    missing_items = []
    warnings = []

    if not profile.company_name:
        missing_items.append("Completare il nome azienda nei dati documento.")
    if not profile.fiscal_note:
        missing_items.append("Completare la nota fiscale nei dati documento.")
    if not estimate.customer.name:
        missing_items.append("Completare il nome cliente.")
    if not estimate.subject:
        missing_items.append("Completare l'oggetto del preventivo.")
    if not estimate.description:
        missing_items.append("Inserire una descrizione della richiesta cliente.")
    if not configuration:
        missing_items.append("Scegliere o aggiungere una configurazione tecnica.")
    elif not configuration.cost_items.exists():
        missing_items.append("Inserire almeno una voce di costo nella configurazione usata.")
    elif not configuration.total:
        missing_items.append("Completare gli importi: il totale della configurazione e ancora zero.")

    if not estimate.valid_until:
        warnings.append("Manca la data di validita della proposta.")
    if not estimate.customer.email:
        warnings.append("Manca l'email cliente.")
    if configuration:
        if not configuration.description:
            warnings.append("Manca la descrizione della configurazione usata.")
        if not configuration.material:
            warnings.append("Manca il materiale o la tecnologia prevista.")

    return {
        "profile": profile,
        "selected_configuration": configuration,
        "missing_items": missing_items,
        "warnings": warnings,
        "can_generate": not missing_items,
    }


def build_consulting_context(estimate):
    checks = build_document_export_checks(estimate)
    configuration = checks["selected_configuration"]
    profile = checks["profile"]
    total = configuration.total if configuration else Decimal("0.00")
    unit_price = configuration.unit_price if configuration else Decimal("0.00")

    return {
        "cliente": {
            "nome": estimate.customer.name,
            "referente": estimate.customer.contact_person or "",
            "email": estimate.customer.email or "",
            "indirizzo": estimate.customer.address or "",
            "codice_fiscale": estimate.customer.tax_code or "",
        },
        "preventivo": {
            "numero": estimate.number,
            "oggetto": estimate.subject,
            "descrizione": estimate.description or "",
            "data": estimate.date.strftime("%d/%m/%Y"),
            "validita": estimate.valid_until.strftime("%d/%m/%Y") if estimate.valid_until else "",
            "quantita": configuration.quantity if configuration else estimate.quantity,
            "condizioni": estimate.general_terms or profile.standard_consulting_terms,
        },
        "configurazione": {
            "nome": configuration.name if configuration else "Configurazione da completare",
            "descrizione": configuration.description if configuration else "",
            "materiale": str(configuration.material) if configuration and configuration.material else "",
            "processo": configuration.process if configuration else "",
            "trattamento": configuration.treatment if configuration else "",
            "durata": configuration.expected_duration if configuration else "",
            "modalita": configuration.operating_mode if configuration else "",
            "note": configuration.public_notes if configuration else "",
            "totale": money(total),
            "unitario": money(unit_price),
        },
        "proposta": {
            "voce": (
                "Compenso per consulenza tecnica, progettazione, validazione "
                "e realizzazione"
            ),
            "totale": money(total),
            "unitario": money(unit_price),
            "nota_fiscale": profile.fiscal_note,
        },
        "b3d": {
            "nome": profile.company_name,
            "sottotitolo": profile.subtitle,
            "indirizzo": profile.address,
            "email": profile.email,
            "telefono": profile.phone,
            "sito": profile.website,
            "codice_fiscale": profile.tax_code,
        },
    }


def build_internal_context(estimate):
    checks = build_document_export_checks(estimate)
    configuration = checks["selected_configuration"]
    profile = checks["profile"]
    cost_items = []
    if configuration:
        cost_items = [
            {
                "categoria": item.get_category_display(),
                "descrizione": item.description,
                "quantita": f"{item.quantity} {item.unit}".strip(),
                "unitario": money(item.unit_cost),
                "totale": money(item.total),
                "note": item.notes,
            }
            for item in configuration.cost_items.all()
        ]

    return {
        **build_consulting_context(estimate),
        "interno": {
            "costo_base": money(configuration.base_cost if configuration else Decimal("0.00")),
            "margine": money(configuration.margin_total if configuration else Decimal("0.00")),
            "margine_percentuale": f"{configuration.margin_percentage:.2f}%" if configuration else "0.00%",
            "note_preventivo": estimate.internal_notes or "",
            "note_configurazione": configuration.internal_notes if configuration else "",
            "nota_footer": profile.internal_footer_note,
            "controlli": checks["warnings"],
        },
        "voci_costo": cost_items,
    }


def build_supply_context(estimate):
    checks = build_document_export_checks(estimate)
    configuration = checks["selected_configuration"]
    profile = checks["profile"]
    supply_items = []
    if configuration:
        supply_items = [
            {
                "categoria": item.get_category_display(),
                "descrizione": item.description,
                "quantita": f"{item.quantity} {item.unit}".strip(),
                "unitario": money(item.unit_cost),
                "totale": money(item.total),
                "note": item.notes,
            }
            for item in configuration.cost_items.filter(visible_supply=True)
        ]

    return {
        **build_consulting_context(estimate),
        "fornitura": {
            "titolo": "Preventivo fornitura / artigiano",
            "voce": "Fornitura e lavorazione additiva",
            "nota_preparatoria": (
                "Template preparatorio per futura modalita fornitura/artigiano. "
                "Dicitura commerciale e fiscale da validare con commercialista prima dell'uso reale."
            ),
            "condizioni": estimate.general_terms or profile.standard_consulting_terms,
            "nota_fiscale": profile.fiscal_note,
        },
        "voci_fornitura": supply_items,
    }


def set_cell_background(cell, color):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell_properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_small_paragraph(cell, text, bold=False, color=None):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_section_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(35, 87, 166)


def create_default_consulting_template(template_path):
    template_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    header = document.add_table(rows=1, cols=2)
    header.autofit = True
    left, right = header.rows[0].cells
    set_cell_background(left, "F3F6FA")
    set_cell_background(right, "2457A6")
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    brand = left.paragraphs[0]
    brand.paragraph_format.space_after = Pt(2)
    run = brand.add_run("{{ b3d.nome }}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 32, 42)
    subtitle = left.add_paragraph()
    run = subtitle.add_run("{{ b3d.sottotitolo }}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(97, 112, 128)
    contact = left.add_paragraph()
    contact.paragraph_format.space_before = Pt(6)
    run = contact.add_run("{{ b3d.email }} | {{ b3d.telefono }} | {{ b3d.sito }}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(97, 112, 128)

    doc_type = right.paragraphs[0]
    doc_type.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = doc_type.add_run("PROPOSTA DI CONSULENZA")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    meta = right.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("N. {{ preventivo.numero }} | {{ preventivo.data }}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    validity = right.add_paragraph()
    validity.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = validity.add_run("Valida fino al {{ preventivo.validita }}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(255, 255, 255)

    document.add_paragraph()

    summary = document.add_table(rows=2, cols=3)
    rows = [
        ("Cliente", "{{ cliente.nome }}"),
        ("Referente", "{{ cliente.referente }}"),
        ("Quantita", "{{ preventivo.quantita }}"),
    ]
    for cell, (label, value) in zip(summary.rows[0].cells, rows):
        set_cell_background(cell, "EEF2F7")
        set_cell_text(cell, label, bold=True)
    for cell, (_, value) in zip(summary.rows[1].cells, rows):
        set_cell_text(cell, value)

    subject = document.add_paragraph()
    subject.paragraph_format.space_before = Pt(10)
    subject.paragraph_format.space_after = Pt(2)
    run = subject.add_run("Oggetto: ")
    run.bold = True
    subject.add_run("{{ preventivo.oggetto }}")

    add_section_title(document, "Contesto e richiesta")
    paragraph = document.add_paragraph("{{ preventivo.descrizione }}")
    paragraph.paragraph_format.space_after = Pt(8)

    add_section_title(document, "Soluzione proposta")
    document.add_paragraph("{{ configurazione.nome }}").runs[0].bold = True
    document.add_paragraph("{{ configurazione.descrizione }}")

    technical = document.add_table(rows=6, cols=2)
    technical.style = "Table Grid"
    rows = [
        ("Materiale / tecnologia", "{{ configurazione.materiale }}"),
        ("Processo", "{{ configurazione.processo }}"),
        ("Trattamento", "{{ configurazione.trattamento }}"),
        ("Durata attesa", "{{ configurazione.durata }}"),
        ("Modalita operativa", "{{ configurazione.modalita }}"),
        ("Note operative", "{{ configurazione.note }}"),
    ]
    for row, (label, value) in zip(technical.rows, rows):
        label_cell, value_cell = row.cells
        set_cell_background(label_cell, "F8FAFC")
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    add_section_title(document, "Sintesi economica")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    labels = ["Voce", "Quantita", "Unitario", "Totale"]
    for cell, label in zip(header, labels):
        set_cell_background(cell, "2457A6")
        set_cell_text(cell, label, bold=True, color="FFFFFF")
    row = table.add_row().cells
    row[0].text = "{{ proposta.voce }}"
    row[1].text = "{{ preventivo.quantita }}"
    row[2].text = "{{ proposta.unitario }}"
    row[3].text = "{{ proposta.totale }}"

    total_row = table.add_row().cells
    total_row[0].merge(total_row[2])
    set_cell_background(total_row[0], "EEF2F7")
    set_cell_background(total_row[3], "EEF2F7")
    set_cell_text(total_row[0], "Totale proposta", bold=True)
    set_cell_text(total_row[3], "{{ proposta.totale }}", bold=True)

    fiscal_note = document.add_paragraph()
    fiscal_note.paragraph_format.space_before = Pt(6)
    run = fiscal_note.add_run("Nota: ")
    run.bold = True
    fiscal_note.add_run("{{ proposta.nota_fiscale }}")

    add_section_title(document, "Condizioni e note")
    document.add_paragraph("{{ preventivo.condizioni }}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("B3D Lab - Documento generato dal gestionale interno")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(97, 112, 128)

    document.save(template_path)
    return template_path


def create_default_internal_template(template_path):
    template_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)

    header = document.add_table(rows=1, cols=2)
    left, right = header.rows[0].cells
    set_cell_background(left, "F3F6FA")
    set_cell_background(right, "17202A")
    set_cell_text(left, "{{ b3d.nome }}", bold=True)
    add_small_paragraph(left, "{{ b3d.sottotitolo }}", color="617080")
    add_small_paragraph(left, "{{ b3d.email }} | {{ b3d.telefono }}", color="617080")
    set_cell_text(right, "SCHEDA INTERNA PREVENTIVO", bold=True, color="FFFFFF")
    marker = right.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = marker.add_run("Non inviare al cliente")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(255, 255, 255)

    add_section_title(document, "Riepilogo")
    summary = document.add_table(rows=7, cols=2)
    summary.style = "Table Grid"
    rows = [
        ("Preventivo", "{{ preventivo.numero }}"),
        ("Cliente", "{{ cliente.nome }}"),
        ("Referente / email", "{{ cliente.referente }} | {{ cliente.email }}"),
        ("Oggetto", "{{ preventivo.oggetto }}"),
        ("Data", "{{ preventivo.data }}"),
        ("Validita", "{{ preventivo.validita }}"),
        ("Configurazione", "{{ configurazione.nome }}"),
    ]
    for row, (label, value) in zip(summary.rows, rows):
        label_cell, value_cell = row.cells
        set_cell_background(label_cell, "EEF2F7")
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    add_section_title(document, "Economia interna")
    totals = document.add_table(rows=2, cols=4)
    rows = [
        ("Costo interno", "{{ interno.costo_base }}"),
        ("Margine", "{{ interno.margine }}"),
        ("Margine percentuale", "{{ interno.margine_percentuale }}"),
        ("Totale proposta", "{{ proposta.totale }}"),
    ]
    for cell, (label, value) in zip(totals.rows[0].cells, rows):
        set_cell_background(cell, "17202A")
        set_cell_text(cell, label, bold=True, color="FFFFFF")
    for cell, (_, value) in zip(totals.rows[1].cells, rows):
        set_cell_background(cell, "F8FAFC")
        set_cell_text(cell, value, bold=True)

    add_section_title(document, "Voci di costo")
    table = document.add_table(rows=4, cols=6)
    table.style = "Table Grid"
    labels = ["Categoria", "Descrizione", "Quantita", "Unitario", "Totale", "Note"]
    for cell, label in zip(table.rows[0].cells, labels):
        set_cell_background(cell, "17202A")
        set_cell_text(cell, label, bold=True, color="FFFFFF")
    table.rows[1].cells[0].text = "{%tr for voce in voci_costo %}"
    cells = table.rows[2].cells
    cells[0].text = "{{ voce.categoria }}"
    cells[1].text = "{{ voce.descrizione }}"
    cells[2].text = "{{ voce.quantita }}"
    cells[3].text = "{{ voce.unitario }}"
    cells[4].text = "{{ voce.totale }}"
    cells[5].text = "{{ voce.note }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"

    add_section_title(document, "Ipotesi operative")
    operative = document.add_table(rows=4, cols=2)
    operative.style = "Table Grid"
    rows = [
        ("Materiale / tecnologia", "{{ configurazione.materiale }}"),
        ("Processo", "{{ configurazione.processo }}"),
        ("Durata attesa", "{{ configurazione.durata }}"),
        ("Modalita operativa", "{{ configurazione.modalita }}"),
    ]
    for row, (label, value) in zip(operative.rows, rows):
        label_cell, value_cell = row.cells
        set_cell_background(label_cell, "F8FAFC")
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    add_section_title(document, "Note interne")
    document.add_paragraph("Preventivo: {{ interno.note_preventivo }}")
    document.add_paragraph("Configurazione: {{ interno.note_configurazione }}")

    add_section_title(document, "Controlli documento")
    document.add_paragraph("{% for controllo in interno.controlli %}{{ controllo }}{% if not loop.last %}\n{% endif %}{% endfor %}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("{{ interno.nota_footer }}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(97, 112, 128)

    document.save(template_path)
    return template_path


def create_default_supply_template(template_path):
    template_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    header = document.add_table(rows=1, cols=2)
    left, right = header.rows[0].cells
    set_cell_background(left, "F3F6FA")
    set_cell_background(right, "3F5F45")
    set_cell_text(left, "{{ b3d.nome }}", bold=True)
    add_small_paragraph(left, "{{ b3d.sottotitolo }}", color="617080")
    add_small_paragraph(left, "{{ b3d.email }} | {{ b3d.telefono }} | {{ b3d.sito }}", color="617080")
    set_cell_text(right, "PREVENTIVO FORNITURA / ARTIGIANO", bold=True, color="FFFFFF")
    meta = right.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run("N. {{ preventivo.numero }} | {{ preventivo.data }}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(255, 255, 255)

    add_section_title(document, "Cliente e richiesta")
    summary = document.add_table(rows=4, cols=2)
    summary.style = "Table Grid"
    rows = [
        ("Cliente", "{{ cliente.nome }}"),
        ("Referente", "{{ cliente.referente }}"),
        ("Oggetto", "{{ preventivo.oggetto }}"),
        ("Quantita", "{{ preventivo.quantita }}"),
    ]
    for row, (label, value) in zip(summary.rows, rows):
        label_cell, value_cell = row.cells
        set_cell_background(label_cell, "EEF2F7")
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    document.add_paragraph("{{ preventivo.descrizione }}")

    add_section_title(document, "Specifiche fornitura")
    technical = document.add_table(rows=5, cols=2)
    technical.style = "Table Grid"
    rows = [
        ("Configurazione", "{{ configurazione.nome }}"),
        ("Materiale / tecnologia", "{{ configurazione.materiale }}"),
        ("Processo", "{{ configurazione.processo }}"),
        ("Trattamento", "{{ configurazione.trattamento }}"),
        ("Note operative", "{{ configurazione.note }}"),
    ]
    for row, (label, value) in zip(technical.rows, rows):
        label_cell, value_cell = row.cells
        set_cell_background(label_cell, "F8FAFC")
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)

    add_section_title(document, "Riepilogo economico")
    table = document.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    labels = ["Voce", "Descrizione", "Quantita", "Unitario", "Totale"]
    for cell, label in zip(table.rows[0].cells, labels):
        set_cell_background(cell, "3F5F45")
        set_cell_text(cell, label, bold=True, color="FFFFFF")
    table.rows[1].cells[0].text = "{%tr for voce in voci_fornitura %}"
    cells = table.rows[2].cells
    cells[0].text = "{{ voce.categoria }}"
    cells[1].text = "{{ voce.descrizione }}"
    cells[2].text = "{{ voce.quantita }}"
    cells[3].text = "{{ voce.unitario }}"
    cells[4].text = "{{ voce.totale }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"

    total_row = table.add_row().cells
    total_row[0].merge(total_row[3])
    set_cell_background(total_row[0], "EEF2F7")
    set_cell_background(total_row[4], "EEF2F7")
    set_cell_text(total_row[0], "Totale fornitura", bold=True)
    set_cell_text(total_row[4], "{{ proposta.totale }}", bold=True)

    add_section_title(document, "Condizioni")
    document.add_paragraph("{{ fornitura.condizioni }}")
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.add_run("Nota fiscale/commerciale: ").bold = True
    note.add_run("{{ fornitura.nota_fiscale }}")
    warning = document.add_paragraph()
    warning.paragraph_format.space_before = Pt(4)
    run = warning.add_run("{{ fornitura.nota_preparatoria }}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(97, 112, 128)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("B3D Lab - Template preparatorio fornitura/artigiano")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(97, 112, 128)

    document.save(template_path)
    return template_path


def get_or_create_default_consulting_template():
    custom_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.CONSULTING_ESTIMATE,
            active=True,
        )
        .exclude(name__startswith="Preventivo consulenza base")
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if custom_template:
        return custom_template

    generated_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.CONSULTING_ESTIMATE,
            active=True,
            name__startswith="Preventivo consulenza base",
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if generated_template:
        relative_path = Path("templates") / "consulting_estimate" / "preventivo_consulenza_base_v3.docx"
        absolute_path = Path(settings.MEDIA_ROOT) / relative_path
        create_default_consulting_template(absolute_path)
        generated_template.template_file = str(relative_path).replace("\\", "/")
        generated_template.template_version = "v3"
        generated_template.notes = (
            "Template base generato dal gestionale e rifinito nello Sprint 11. "
            "Resta sostituibile con un template DOCX personalizzato."
        )
        generated_template.save(update_fields=["template_file", "template_version", "notes"])
        return generated_template

    template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.CONSULTING_ESTIMATE,
            active=True,
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if template:
        return template

    relative_path = Path("templates") / "consulting_estimate" / "preventivo_consulenza_base_v3.docx"
    absolute_path = Path(settings.MEDIA_ROOT) / relative_path
    create_default_consulting_template(absolute_path)

    return DocumentTemplate.objects.create(
        name="Preventivo consulenza base",
        document_type=DocumentTemplate.DocumentType.CONSULTING_ESTIMATE,
        profile="consulting",
        template_file=str(relative_path).replace("\\", "/"),
        template_version="v3",
        active=True,
        notes=(
            "Template base generato dal gestionale e rifinito nello Sprint 11. "
            "Resta sostituibile con un template DOCX personalizzato."
        ),
    )


def get_or_create_default_internal_template():
    custom_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.INTERNAL_ESTIMATE,
            active=True,
        )
        .exclude(name__startswith="Preventivo interno base")
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if custom_template:
        return custom_template

    generated_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.INTERNAL_ESTIMATE,
            active=True,
            name__startswith="Preventivo interno base",
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    relative_path = Path("templates") / "internal_estimate" / "preventivo_interno_base_v2.docx"
    absolute_path = Path(settings.MEDIA_ROOT) / relative_path
    create_default_internal_template(absolute_path)
    if generated_template:
        generated_template.template_file = str(relative_path).replace("\\", "/")
        generated_template.template_version = "v2"
        generated_template.notes = (
            "Template interno generato dal gestionale e rifinito nello Sprint 11. "
            "Mostra riepilogo, costi, margine, ipotesi operative, note e controlli."
        )
        generated_template.save(update_fields=["template_file", "template_version", "notes"])
        return generated_template

    template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.INTERNAL_ESTIMATE,
            active=True,
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if template:
        return template

    return DocumentTemplate.objects.create(
        name="Preventivo interno base",
        document_type=DocumentTemplate.DocumentType.INTERNAL_ESTIMATE,
        profile="internal",
        template_file=str(relative_path).replace("\\", "/"),
        template_version="v2",
        active=True,
        notes=(
            "Template interno generato dal gestionale e rifinito nello Sprint 11. "
            "Mostra riepilogo, costi, margine, ipotesi operative, note e controlli."
        ),
    )


def get_or_create_default_supply_template():
    custom_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.SUPPLY_ESTIMATE,
            active=True,
        )
        .exclude(name__startswith="Preventivo fornitura base")
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if custom_template:
        return custom_template

    generated_template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.SUPPLY_ESTIMATE,
            active=True,
            name__startswith="Preventivo fornitura base",
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    relative_path = Path("templates") / "supply_estimate" / "preventivo_fornitura_base_v1.docx"
    absolute_path = Path(settings.MEDIA_ROOT) / relative_path
    create_default_supply_template(absolute_path)
    if generated_template:
        generated_template.template_file = str(relative_path).replace("\\", "/")
        generated_template.template_version = "v1"
        generated_template.notes = (
            "Template preparatorio fornitura/artigiano generato nello Sprint 14. "
            "Da validare commercialmente e fiscalmente prima dell'uso reale."
        )
        generated_template.save(update_fields=["template_file", "template_version", "notes"])
        return generated_template

    template = (
        DocumentTemplate.objects.filter(
            document_type=DocumentTemplate.DocumentType.SUPPLY_ESTIMATE,
            active=True,
        )
        .order_by("-uploaded_at", "-id")
        .first()
    )
    if template:
        return template

    return DocumentTemplate.objects.create(
        name="Preventivo fornitura base",
        document_type=DocumentTemplate.DocumentType.SUPPLY_ESTIMATE,
        profile="supply",
        template_file=str(relative_path).replace("\\", "/"),
        template_version="v1",
        active=True,
        notes=(
            "Template preparatorio fornitura/artigiano generato nello Sprint 14. "
            "Da validare commercialmente e fiscalmente prima dell'uso reale."
        ),
    )


def generate_estimate_docx(estimate, document_type, template, context, folder, filename_label, notes):
    for _ in range(3):
        latest_version = (
            GeneratedDocument.objects.filter(
                estimate=estimate,
                document_type=document_type,
            ).aggregate(value=Max("version"))["value"]
            or 0
        )
        version = latest_version + 1
        filename = f"preventivo_{safe_filename(estimate.number)}_{filename_label}_v{version}.docx"
        relative_path = Path("generated") / safe_filename(estimate.number) / folder / filename
        output_path = Path(settings.MEDIA_ROOT) / relative_path

        render_docx_template(
            template.template_file.path,
            output_path,
            context,
        )
        pdf_path = convert_docx_to_pdf(output_path, output_path.parent)
        pdf_file = ""
        if pdf_path:
            pdf_file = str(relative_path.with_suffix(".pdf")).replace("\\", "/")

        try:
            with transaction.atomic():
                return GeneratedDocument.objects.create(
                    estimate=estimate,
                    template=template,
                    document_type=document_type,
                    version=version,
                    docx_file=str(relative_path).replace("\\", "/"),
                    pdf_file=pdf_file,
                    notes=notes,
                )
        except IntegrityError:
            continue

    raise IntegrityError("Impossibile assegnare una nuova versione al documento generato.")


def generate_consulting_estimate_docx(estimate):
    template = get_or_create_default_consulting_template()
    return generate_estimate_docx(
        estimate=estimate,
        document_type=GeneratedDocument.DocumentType.CONSULTING,
        template=template,
        context=build_consulting_context(estimate),
        folder="consulting",
        filename_label="consulenza",
        notes="Documento cliente consulenza generato dal dettaglio preventivo.",
    )


def generate_internal_estimate_docx(estimate):
    template = get_or_create_default_internal_template()
    return generate_estimate_docx(
        estimate=estimate,
        document_type=GeneratedDocument.DocumentType.INTERNAL,
        template=template,
        context=build_internal_context(estimate),
        folder="internal",
        filename_label="interno",
        notes="Documento interno dettagliato generato dal dettaglio preventivo.",
    )


def generate_supply_estimate_docx(estimate):
    template = get_or_create_default_supply_template()
    return generate_estimate_docx(
        estimate=estimate,
        document_type=GeneratedDocument.DocumentType.SUPPLY,
        template=template,
        context=build_supply_context(estimate),
        folder="supply",
        filename_label="fornitura",
        notes=(
            "Documento fornitura/artigiano preparatorio generato dal dettaglio preventivo. "
            "Da validare prima dell'uso reale verso cliente."
        ),
    )
